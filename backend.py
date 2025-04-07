# ======================================================================================================
#                                               imports        
# ======================================================================================================

from openai import OpenAI
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from pydantic import BaseModel, Field,HttpUrl
import re
import datetime
import logging
import os
import traceback
from typing import Optional, List, Dict
from datetime import datetime, timedelta
import pandas as pd
import plotly.express as px
import calendar
import pymongo
from pdf_maker import pdfmaker
import base64
import time
import json
from docx import Document
from pypdf import PdfReader

# ======================================================================================================
#                                       logging configuration       
# ======================================================================================================

# Create logs directory if it doesn't exist
os.makedirs("logs", exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        # logging.FileHandler("logs/app.log"),
        logging.StreamHandler()
    ]
)

# Create specialized loggers
ollama_logger = logging.getLogger('ollama')
resume_logger = logging.getLogger('resume')
meeting_logger = logging.getLogger('meeting')
streamlit_logger = logging.getLogger('streamlit')

# ======================================================================================================
#                                       ollama llm connect        
# ======================================================================================================

ollama_client = OpenAI(
    base_url="http://localhost:11434/v1",
    api_key="ollama"
)

ollama_logger.info("Initializing Ollama client with base URL: http://localhost:11434/v1")

model = "llama3.2"
# model = "mistral"
ollama_logger.info(f"Using model: {model}")

message = [
    {
        "role": "system",
        "content": """You are a highly accurate and comprehensive resume analyzer AI assistant. Your task is to meticulously extract and analyze resume information with extreme precision. Follow these guidelines:

1. Personal Details: Extract full name, contact information (email, phone), online profiles (LinkedIn, GitHub), and complete address information.
2. Professional Summary: Identify and rewrite the professional summary to highlight key qualifications, career focus, and unique value proposition.
3. Work Experience: Carefully analyze employment history, identifying job titles, companies, dates, responsibilities, and quantifiable achievements. Pay special attention to action verbs, metrics, and impact statements.
4. Education: Extract all educational credentials including degrees, institutions, years, specializations, GPA/percentages, and relevant coursework.
5. Skills: Comprehensively identify technical skills, soft skills, programming languages, tools, and methodologies. Categorize skills by proficiency level when possible.
6. Projects: Extract all projects with detailed descriptions, technologies used, role, impact, and outcomes.
7. Certifications: Identify all professional certifications with issuing organizations and dates.
8. Additional Information: Extract hobbies, interests, languages, volunteering experience, and other relevant personal information.
9. Achievements: Identify awards, publications, patents, or other notable accomplishments with dates and significance.

When information is not explicitly stated but strongly implied, note this in your extraction. For unclear or ambiguous information, mark it as "Requires verification" rather than making assumptions.
"""
    },
]

ollama_logger.info("Setting up system message for resume analysis")

try:
    response = ollama_client.chat.completions.create(
        model=model,
        messages=message,
    )
    ollama_logger.info("Initial Ollama connection successful")
except Exception as e:
    ollama_logger.error(f"Failed to connect to Ollama: {str(e)}")

# ======================================================================================================
#                                   Pydantic response format for RESUME
# ======================================================================================================

resume_logger.info("Defining Pydantic models for resume parsing")


class PersonalDetails(BaseModel):
    """Personal Details: Comprehensive identification of resume candidate"""
    Full_Name: str = Field(description="Full name of resume candidate")
    Email_Address: str = Field(description="Email address of resume candidate")
    Phone_Number: Optional[str] = Field(description="Phone number of resume candidate") # Changed to string to handle international formats
    LinkedIn_Profile: Optional[str] = Field(description="LinkedIn profile URL or username")
    GitHub_Profile: Optional[str] = Field(description="GitHub profile URL or username")
    Portfolio_Website: Optional[str] = Field(description="Personal or portfolio website URL")
    Address: Optional[str] = Field(description="Address of resume candidate")
    City: Optional[str] = Field(description="City where candidate resides")
    State: Optional[str] = Field(description="State or province where candidate resides")
    Country: Optional[str] = Field(description="Country where candidate resides")
    Pincode: Optional[str] = Field(description="Postal/Zip code of address") # Changed to string to handle various formats

class ProfessionalSummary(BaseModel):
    """Professional Summary: Detailed overview of professional profile"""
    Summary: str = Field(description="Professional summary highlighting experience, skills and achievements")
    Objective: Optional[str] = Field(description="Career objective or professional goals")
    Years_of_Experience: Optional[int] = Field(description="Total years of professional experience")
    Industry_Focus: Optional[List[str]] = Field(description="Primary industries or sectors of expertise")

class ExperienceItem(BaseModel):
    company: str
    title: str
    location: Optional[str]
    duration: str
    start_date: Optional[str] = Field(description="Start date of employment")
    end_date: Optional[str] = Field(description="End date of employment or 'Present' if current")
    technologies_used: Optional[List[str]] = Field(description="Technologies or tools used in this role")

class ExperienceList(BaseModel):
    list_of_experience: List[ExperienceItem]

class EducationItem(BaseModel):
    degree: str
    institution: str
    years: str
    start_date: Optional[str] = Field(description="Start date of education")
    end_date: Optional[str] = Field(description="End date of education or 'Present' if current")
    percentage: Optional[float] = Field(description="GPA, percentage or other academic score")
    specialization: Optional[str] = Field(description="Major, specialization or focus area")
    relevant_coursework: Optional[List[str]] = Field(description="Relevant courses completed")
    achievements: Optional[List[str]] = Field(description="Academic achievements or honors")

class EducationList(BaseModel):
    list_of_education: List[EducationItem]

class SkillsDetails(BaseModel):
    """Skills: Comprehensive categorization of candidate's capabilities"""
    Technical_skills: List[str] = Field(description="List of technical skills")
    Soft_Skills: Optional[List[str]] = Field(description="List of soft/interpersonal skills")
    Programming_Languages: Optional[List[str]] = Field(description="Programming languages with proficiency level when available")
    Frameworks_Libraries: Optional[List[str]] = Field(description="Frameworks and libraries the candidate is familiar with")
    Tools_Software: Optional[List[str]] = Field(description="Tools, software or platforms the candidate has used")
    Methodologies: Optional[List[str]] = Field(description="Methodologies, processes or approaches the candidate is familiar with")
    Languages: Optional[List[str]] = Field(description="Human languages spoken with proficiency level")

class CertificationsDetails(BaseModel):
    """Certifications: Professional qualifications and trainings"""
    Certification_name: str = Field(description="Name of the certification or course")
    Issuing_organization: str = Field(description="Organization that issued the certification")

class CertificationsList(BaseModel):
    list_of_certificates: List[CertificationsDetails]

class ProjectsDetails(BaseModel):
    """Projects: Detailed analysis of project experience"""
    Project_name: str = Field(description="Project name")
    Project_description: str = Field(description="Comprehensive project description")
    Role: Optional[str] = Field(description="Role in the project")
    Duration: Optional[str] = Field(description="Duration of project involvement")
    Technologies_used: Optional[List[str]] = Field(description="Technologies, tools or methods used")
    Team_size: Optional[int] = Field(description="Size of the project team")
    URL: Optional[str] = Field(description="Project URL or repository link")
    Key_achievements: Optional[List[str]] = Field(description="Notable achievements or outcomes")

class ProjectList(BaseModel):
    list_of_projects: List[ProjectsDetails]

class Additional_Information(BaseModel):
    Hobbies: Optional[List[str]] = Field(description="List of hobbies")
    Interests: Optional[List[str]] = Field(description="List of professional or personal interests")
    Languages: Optional[List[str]] = Field(description="Languages spoken with proficiency level")
    Availability: Optional[str] = Field(description="Availability for work, notice period, etc.")

class AchievementsDetails(BaseModel):
    """Achievements: Recognition and notable accomplishments"""
    Achievement_description: str = Field(description="Description of the achievement")
    Date: Optional[str] = Field(description="Date of achievement")
    Awarding_organization: Optional[str] = Field(description="Organization that gave the award/recognition")
    Impact: Optional[str] = Field(description="Impact or significance of the achievement")

class AchievementsList(BaseModel):
    list_of_achievements: List[AchievementsDetails]

# ======================================================================================================
#                                   Pydantic response format for Meeting
# ======================================================================================================

meeting_logger.info("Defining Pydantic models for meeting extraction")

class MeetingInfo(BaseModel):
    """Extracted information about whether the user input relates to a meeting."""
    meeting_description: str = Field(description="Brief description of the meeting purpose.")
    is_calendar_event: bool = Field(description="Indicates if the input is a calendar event.")
    confidence_score: float = Field(description="Confidence level (0 to 1) for meeting detection.")
    title: str = Field(description="The title or subject of the meeting.")
    datetime: str = Field(description="The date and time of the meeting.format example: 2025-03-26T14:00:00")
    date: str = Field(description="The date of the meeting in match format %A, %yyyy-%mm-%dd.")
    time: str = Field(description="The time of the meeting in match format  %HH:%MM:%SS.")
    duration_of_meeting: float = Field(description="Duration of the meeting in hours.")
    participants: List[str] = Field(description="List of participants attending the meeting.")
    confirmation_message: str = Field(description="Natural language confirmation message for the user.")
    calendar_link: Optional[str] = Field(description="Generated calendar link, if applicable.")
    notes: Optional[str] = Field(description="Additional notes or details about the meeting.")
    location: Optional[str] = Field(description="Location of the meeting.")



# ======================================================================================================
#                                   Pydantic response format for classifier
# ======================================================================================================

meeting_logger.info("Defining Pydantic models for classifier extraction")

class ClassifierInfo(BaseModel):
    """Extracted information about the classification of the resume."""
    category: str = Field(description="The classified category of the resume.")
    confidence_score: float = Field(description="Confidence level (0 to 1) for classification.")
    description: Optional[str] = Field(description="Additional details or reasoning for the classification.")


# ======================================================================================================
#                                     similarity search prompts        
# ======================================================================================================

resume_logger.info("Setting up similarity search prompts for resume sections")

msg1 = """Extract all personal details from the resume with high precision, including:
- Full name (exact spelling and format)
- Email address (complete and accurate)
- Phone number (with correct formatting)
- LinkedIn profile URL or handle
- GitHub profile URL or username
- Portfolio website if available
- Complete address details including city, state, country, and postal/zip code
Look for these details typically found at the top of the resume or in a dedicated contact section. If certain information is not present, indicate this rather than making assumptions."""

msg2 = """Extract the professional summary section with careful attention to detail:
- Main professional summary paragraph highlighting experience and skills
- Career objective or professional goals statement if present
- Key qualifications or core competencies listed
- Implied years of total experience (look for phrases like "X years of experience")
- Primary industries or sectors where the candidate has worked
Pay attention to the language used to describe the candidate's value proposition, career trajectory, and professional identity."""

msg3 = """Extract comprehensive work experience details with precision:
- Job titles (exact titles as listed)
- Company names (full legal names when provided)
- Location information for each position
- Detailed responsibilities (rewritten as full sentences starting with action verbs)
- Projects or initiatives led or contributed to
- Technologies or methodologies used in each role
Look for chronological work history, typically in reverse chronological order. Pay special attention to promotions within the same company."""

msg4 = """Extract detailed education information with precision:
- Full degree names (e.g., "Bachelor of Science in Computer Science" rather than just "BS")
- Complete institution names
- Years of attendance (both start and end dates)
- GPA, percentage, or other academic scores if provided
- Major, specialization, concentration or focus areas
- Relevant coursework, thesis topics, or research projects
- Academic honors, scholarships, or distinctions
- Study abroad experiences if mentioned
Focus on higher education first, followed by professional education or training."""

msg5 = """Extract and categorize all skills listed on the resume with detailed organization:
- Technical skills (software, methods, techniques)
- Soft/interpersonal skills
- Programming languages (with proficiency levels if indicated)
- Frameworks and libraries
- Tools, platforms and software applications
- Methodologies and processes (e.g., Agile, Six Sigma)
- Human languages spoken (with proficiency levels)
Look for skills mentioned throughout the resume, not just in a dedicated skills section. Skills may be implied in project descriptions or work responsibilities."""

msg6 = """Extract comprehensive certification information:
- Complete certification names (exact titles)
- Full names of issuing organizations or institutions
- Dates obtained and expiration dates if provided
- Credential IDs or verification numbers if listed
- Areas of specialization within certifications
Look for both formal certifications and completed training programs or courses."""

msg7 = """Extract detailed project information with high precision:
- Project names and complete descriptions
- Your specific role or contributions
- Project duration and timeline
- Team size and your leadership responsibilities if applicable
- Technologies, tools, and methodologies used
- Key challenges addressed and solutions implemented
- Quantifiable outcomes, impacts, or results
- URLs or repositories if mentioned
Look for projects mentioned throughout the resume, including those within work experience sections, education sections, or dedicated project sections."""

msg8 = """Extract comprehensive additional information that provides context about the candidate:
- Hobbies and personal interests
- Professional interests or areas of passion
- Volunteer experience or community involvement
- Languages spoken (with proficiency levels)
- Availability or notice period information
- Relocation preferences or work arrangement preferences
This information is typically found in sections like "Additional Information," "Personal Interests," or at the end of the resume."""

msg9 = """Extract all achievements, awards, honors, and publications with detailed precision:
- Complete descriptions of achievements or awards
- Exact names of awarding organizations or publications
- Dates received or published
- The significance or impact of each achievement
- Selection criteria or competition details if mentioned
- For publications: co-authors, journals, conferences, or publishing venues
Look for achievements mentioned throughout the resume, including within work experience, education, or dedicated sections."""

msg10="""
You are a highly accurate and comprehensive resume title classifier AI assistant. Your task is to classify the resume into one of the following categories based on given json data.

Computer Science / IT / Software Engineering
Data Science / AI / Machine Learning
Electrical Engineering
Mechanical Engineering
Civil Engineering
Marketing
Finance
Human Resources (HR)
Sales
Business Administration
Legal / Law
Graphic Design / UX/UI
Biomedical Science
Education / Teaching
Psychology
"""

msg = [msg1, msg2, msg3, msg4, msg5, msg6, msg7, msg8, msg9, msg10]

# ======================================================================================================
#                             get llm response using system prompt and classes for RESUME     
# ======================================================================================================
def improved_text_splitter(text: str) -> List[str]:
    """Advanced semantic-aware text splitting for resumes with better section recognition"""
    # Expanded section pattern recognition
    section_patterns = [
        r'\n\n(?:INTRODUCTION|ABOUT ME|PROFILE|SUMMARY|PROFESSIONAL SUMMARY)\n\n',
        r'\n\n(?:WORK EXPERIENCE|PROFESSIONAL EXPERIENCE|EMPLOYMENT HISTORY|WORK HISTORY)\n\n',
        r'\n\n(?:EDUCATION|ACADEMIC BACKGROUND|EDUCATIONAL QUALIFICATIONS|ACADEMIC QUALIFICATIONS)\n\n',
        r'\n\n(?:SKILLS|TECHNICAL SKILLS|CORE COMPETENCIES|KEY SKILLS|PROFESSIONAL SKILLS)\n\n',
        r'\n\n(?:PROJECTS|KEY PROJECTS|PROJECT EXPERIENCE|TECHNICAL PROJECTS)\n\n',
        r'\n\n(?:CERTIFICATIONS|CERTIFICATES|COURSES|TRAINING|PROFESSIONAL DEVELOPMENT)\n\n',
        r'\n\n(?:PERSONAL INFORMATION|PERSONAL DETAILS|CONTACT INFORMATION|CONTACT DETAILS)\n\n',
        r'\n\n(?:HOBBIES|INTERESTS|ACTIVITIES|EXTRACURRICULAR ACTIVITIES)\n\n',
        r'\n\n(?:ACHIEVEMENTS|AWARDS|HONORS|ACCOMPLISHMENTS|PUBLICATIONS)\n\n',
        r'\n\n(?:ADDITIONAL INFORMATION|OTHER INFORMATION|MISCELLANEOUS)\n\n',
        r'\n\n(?:LANGUAGES|LANGUAGE PROFICIENCY)\n\n',
        r'\n\n(?:VOLUNTEER EXPERIENCE|COMMUNITY SERVICE)\n\n',
        r'\n\n(?:REFERENCES|PROFESSIONAL REFERENCES)\n\n'
    ]
    
    # Find all section boundaries
    section_boundaries = []
    for pattern in section_patterns:
        for match in re.finditer(pattern, text, re.IGNORECASE):
            section_boundaries.append((match.start(), match.end(), match.group().strip()))
    
    # Sort boundaries by position
    section_boundaries.sort(key=lambda x: x[0])
    
    # Extract sections based on boundaries
    sections = []
    if not section_boundaries:
        # If no sections found, treat the entire text as one section
        sections.append(text)
    else:
        # Add text before first section if any
        if section_boundaries[0][0] > 0:
            sections.append(text[:section_boundaries[0][0]])
        
        # Process each section
        for i, (start, end, header) in enumerate(section_boundaries):
            # Find end of current section (start of next section or end of text)
            section_end = section_boundaries[i+1][0] if i < len(section_boundaries)-1 else len(text)
            sections.append(header + text[end:section_end])
    
    # Use more sophisticated content splitter with better preservation of context
    content_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1100,  # Smaller chunk size for more precise retrieval
        chunk_overlap=250,  # Larger overlap to maintain context
        length_function=len,
        separators=[
            "\n\n\n",  # Major section breaks
            "\n\n",     # Paragraph breaks
            "\n• ", "\n- ", "\n* ", "\n○ ",  # List items
            ". ", "! ", "? ",  # Sentence boundaries
            ", ", "; ", " ", ""  # Word boundaries
        ]
    )
    
    final_chunks = []
    for section in sections:
        # Preserve section context in each chunk
        section_match = re.match(r'^(\n\n[A-Z ]+\n\n)', section)
        section_header = section_match.group(0) if section_match else ""
        
        if section_header:
            section_content = section[len(section_header):]
            # Split the content
            sub_chunks = content_splitter.split_text(section_content)
            # Add header to each sub-chunk
            for chunk in sub_chunks:
                final_chunks.append(section_header + chunk)
        else:
            chunks = content_splitter.split_text(section)
            final_chunks.extend(chunks)
    
    return final_chunks

def clean_resume_text(text: str) -> str:
    """Enhanced text cleaning for resumes with better preprocessing and normalization"""
    # Convert to plain text if needed (remove any HTML tags)
    text = re.sub(r'<[^>]+>', ' ', text)
    
    # Standardize newlines and whitespace
    text = re.sub(r'\r\n|\r', '\n', text)  # Standardize line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)  # Replace excessive newlines
    text = re.sub(r'[ \t]+', ' ', text)     # Standardize spaces and tabs
    
    # Remove control characters and non-printable characters
    text = re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)
    
    # Remove common resume artifacts, headers, and footers
    text = re.sub(r'(?i)\b(page\s*\d+|\d+\s*of\s*\d+|confidential|resume|cv|curriculum vitae)\b', '', text)
    text = re.sub(r'[-_=]{3,}', '', text)  # Remove separator lines
    
    # Fix common OCR and formatting issues
    text = re.sub(r'•', '- ', text)  # Standardize bullet points
    text = re.sub(r'(\d+)\.(\d+)', r'\1.\2', text)  # Fix merged decimals
    text = re.sub(r'(\d+)/(\d+)/(\d+)', r'\1/\2/\3', text)  # Fix merged dates
    
    # Standardize common section headers with comprehensive patterns
    section_headers = {
        r'(?i)\b(introduction|about\s*me|profile|summary|professional\s*summary)\b': 'INTRODUCTION',
        r'(?i)\b(work\s*history|employment\s*history|professional\s*experience|work\s*experience)\b': 'WORK EXPERIENCE',
        r'(?i)\b(education|academic\s*background|educational\s*qualifications)\b': 'EDUCATION',
        r'(?i)\b(skills|technical\s*skills|core\s*competencies|key\s*skills)\b': 'SKILLS',
        r'(?i)\b(personal\s*details?|personal\s*information|contact\s*details?|contact\s*information)\b': 'PERSONAL INFORMATION',
        r'(?i)\b(projects|key\s*projects|project\s*experience|technical\s*projects)\b': 'PROJECTS',
        r'(?i)\b(certifications|certificates|professional\s*development|courses|training)\b': 'CERTIFICATIONS',
        r'(?i)\b(hobbies|interests|activities|extracurricular\s*activities)\b': 'HOBBIES',
        r'(?i)\b(achievements|awards|honors|accomplishments|publications)\b': 'ACHIEVEMENTS',
        r'(?i)\b(additional\s*information|other\s*information|miscellaneous)\b': 'ADDITIONAL INFORMATION',
        r'(?i)\b(languages|language\s*proficiency)\b': 'LANGUAGES',
        r'(?i)\b(volunteer\s*experience|community\s*service)\b': 'VOLUNTEER EXPERIENCE',
        r'(?i)\b(references|professional\s*references)\b': 'REFERENCES',
    }
    
    for pattern, replacement in section_headers.items():
        text = re.sub(pattern, f'\n\n{replacement}\n\n', text)
    
    return text.strip()

def get_llm_response(userInput: str, message: str, response_format: type[BaseModel]) -> BaseModel:
        try:
            resume_logger.info(f"Extracting {response_format.__name__} details")
            
            # Enhanced error handling with retries
            max_retries = 2
            retry_count = 0
            
            while retry_count < max_retries:
                try:
                    # Add temperature for more deterministic outputs
                    response = ollama_client.beta.chat.completions.parse(
                        model=model,
                        messages=[
                            {"role": "system", "content": message},
                            {"role": "user", "content": userInput} 
                        ],
                        response_format=response_format,
                        temperature=0.1 
                    )
                    resume_logger.info(f"Successfully extracted {response_format.__name__} details")
                    return response.choices[0].message.parsed
                except Exception as e:
                    retry_count += 1
                    resume_logger.warning(f"Attempt {retry_count} failed: {str(e)}")
                    if retry_count >= max_retries:
                        raise
                    
        except Exception as e:
            resume_logger.error(f"Error getting LLM response: {str(e)}")
            raise

def fn_Resume(resume_text: str) -> dict:
    resume_logger.info("Starting full resume analysis")
    
    if not resume_text.strip():
        resume_logger.warning("Empty resume text provided")
        return {}

    try:
        # Better text cleaning and preprocessing
        cleaned_text = clean_resume_text(resume_text)
        
        # Improved semantic chunking
        documents = improved_text_splitter(cleaned_text)
        resume_logger.info(f"Split resume into {len(documents)} semantic chunks")

        # Use a more sophisticated embedding model if available
        try:
            embedding_model = "nomic-embed-text:latest"
            resume_logger.info(f"Using embedding model: {embedding_model}")
            embeddings = OllamaEmbeddings(model=embedding_model)
        except Exception as e:
            resume_logger.warning(f"Could not use specified embedding model, falling back to default: {str(e)}")
            embeddings = OllamaEmbeddings(model="nomic-embed-text")

        # Use a dummy embedding if no text is available
        if not documents:
            resume_logger.warning("No document chunks were created, using placeholder text")
            db = FAISS.from_texts(["Resume appears to be empty or unreadable"], embeddings)
        else:
            resume_logger.info("Creating FAISS vector store for semantic search")
            db = FAISS.from_texts(documents, embeddings)

        # Enhanced retrieval with more context per section
        results = []
        for i, query in enumerate(msg):
            resume_logger.info(f"Performing similarity search for section {i+1}")
            # Retrieve more documents for complex sections
            k_values = {
                0: 5,  # Personal Details
                1: 5,  # Professional Summary 
                2: 5,  # Work Experience (more context needed)
                3: 5,  # Education
                4: 7,  # Skills (spread throughout document)
                5: 4,  # Certifications
                6: 6,  # Projects (may be scattered)
                7: 4,  # Hobbies
                8: 4   # Achievements
            }
            
            rs = db.similarity_search(query, k=k_values.get(i, 5))
            tx = " ".join([doc.page_content for doc in rs])
            results.append(tx)
            resume_logger.debug(f"Section {i+1} search results length: {len(tx)} characters")
        
        # Extract sections with enhanced error handling
        sections = {}
        section_models = [
            ("Personal_Details", results[0], msg1, PersonalDetails),
            ("Professional_Summary", results[1], msg2, ProfessionalSummary),
            ("Work_Experience", results[2], msg3, ExperienceList),
            ("Education_Details", results[3], msg4, EducationList),
            ("Skills_Details", results[4], msg5, SkillsDetails),
            ("Certifications_Details", results[5], msg6, CertificationsList),
            ("Projects_Details", results[6], msg7, ProjectList),
            ("Additional_Information", results[7], msg8, Additional_Information),
            ("Achievements_Details", results[8], msg9, AchievementsList),
        ]
        
        for section_name, result_text, prompt, model_class in section_models:
            try:
                sections[section_name] = get_llm_response(result_text, prompt, model_class)
                resume_logger.info(f"Successfully processed {section_name}")
            except Exception as e:
                resume_logger.error(f"Error processing {section_name}: {str(e)}")
                sections[section_name] = f"Error extracting {section_name}: {str(e)}"

        try:
            sections['Classification'] = fn_classifier(
                {
                    'Summary': sections['Professional_Summary'].Summary,
                    'Work_Experience': [exp.model_dump() for exp in sections['Work_Experience'].list_of_experience],
                    'Education_Details': [edu.model_dump() for edu in sections['Education_Details'].list_of_education],
                    'Technical_skills': sections['Skills_Details'].Technical_skills
                }
            )
        except Exception as e:
            resume_logger.error(f"Error during classification: {str(e)}")
            sections['Classification'] = f"Error during classification: {str(e)}"
        
        # Add metadata about the extraction process
        sections["metadata"] = {
            "extraction_timestamp": datetime.now().isoformat(),
            "model_used": model,
            "chunks_processed": len(documents),
            "version": "2.0.0"
        }
        
        return sections

    except Exception as e:
        resume_logger.error(f"Critical error in complete resume analysis: {str(e)}")
        traceback.print_exc()
        return {"error": str(e), "traceback": traceback.format_exc()}

# ======================================================================================================
#                          get llm response using system prompt and classes for Meeting
# ======================================================================================================

def extract_meeting_info(user_input: str) -> Optional[MeetingInfo]:
    """Extracts structured meeting information using FAISS and LLM."""
    meeting_logger.info("Starting meeting extraction process")
    
    try:
        meeting_logger.info("Calling Ollama model for meeting extraction")
        response = ollama_client.beta.chat.completions.parse(
            model=model,
            messages=[
                {
                    'role': 'system',
                    'content': '''
                        You are a meeting information extractor. Your task is to find meeting details in user text and put them into a specific JSON format.
                        Follow these instructions strictly and return only the JSON.

                        Expected JSON Format:
                        {
                            "meeting_description": "Brief description of the meeting purpose.",
                            "is_calendar_event": true/false,
                            "confidence_score": float (0 to 1),
                            "title": "The title or subject of the meeting.",
                            "datetime": "YYYY-MM-DDTHH:MM:SS",
                            "date": "Day, YYYY-MM-DD",
                            "time": "HH:MM:SS",
                            "duration_of_meeting": float (in hours),
                            "participants": ["List", "of", "participants"],
                            "confirmation_message": "Natural language confirmation message for the user.",
                            "calendar_link": "Generated calendar link if applicable, else null."
                            "notes": "Additional notes or details about the meeting."
                            "location":"Location of the meeting. like 'Google Meet', 'Zoom', 'Aether Office', etc."
                        }

                        Ensure accuracy in parsing and extracting details. If any field is missing or unclear, make a reasonable assumption.
                    '''
                },
                {'role': 'user', 'content': user_input}
            ],
            response_format=MeetingInfo
        )

        info = response.choices[0].message.parsed
        meeting_logger.info(f"LLM extraction successful - Title: '{info.title}', Date: '{info.date}', Time: '{info.time}'")

        if not info or not info.is_calendar_event or info.confidence_score < 0.6:
            meeting_logger.warning("No valid meeting detected or confidence too low")
            print("❌ No valid meeting detected.")
            return None
        
        meeting_logger.info("Meeting extraction completed successfully")
        # Final aggregated data
        return {
            "meeting_description":info.meeting_description,
            "is_calendar_event":info.is_calendar_event,
            "confidence_score":info.confidence_score,
            "title":info.title,
            "datetime":info.datetime,
            "date":info.date,
            "time":info.time,
            "duration_of_meeting":info.duration_of_meeting,
            "participants":info.participants,
            "confirmation_message":info.confirmation_message,
            "calendar_link":info.calendar_link,
            "notes":info.notes,
            "location":info.location
        }
    except Exception as e:
        meeting_logger.error(f"Error in meeting extraction process: {str(e)}")
        raise

# ======================================================================================================
#                          get llm response using system prompt and classes for title
# ======================================================================================================

    
def fn_classifier(all_d: dict) -> str:
    """Generate a title for the resume based on its content."""
    resume_logger.info("Generating title classifier for user data")
    
    try:
        # Convert the dictionary to a JSON string
        all_d_str = json.dumps(all_d)
        
        # Call the LLM with the prompt
        response = ollama_client.beta.chat.completions.parse(
            model=model,
            messages=[
                {
                    "role": "system", 
                    "content": '''
                            You are a highly accurate and comprehensive resume title classifier AI assistant. Your task is to classify the resume into one of the following categories based on given json data.

                            Computer Science / IT / Software Engineering
                            Data Science / AI / Machine Learning
                            Electrical Engineering
                            Mechanical Engineering
                            Civil Engineering
                            Marketing
                            Finance
                            Human Resources (HR)
                            Sales
                            Business Administration
                            Legal / Law
                            Graphic Design / UX/UI
                            Biomedical Science
                            Education / Teaching
                            Psychology
                '''},
                {"role": "user", "content": all_d_str}
            ],
            response_format=ClassifierInfo
        )
        
        resume_logger.info(f"Generated title: {response.choices[0].message.parsed.category}")
        
        return response.choices[0].message.parsed
    except Exception as e:
        resume_logger.error(f"Error generating title: {str(e)}")
        raise



# ======================================================================================================
#                                       DATA STORAGE CONFIGURATION in MongoDB
# ======================================================================================================

MEETINGS_FILE = "meeting_records"
RESUMES_FILE = "job_application"

myclient = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = myclient["resumedb"]

def load_data(collection_name: str) -> List[Dict]:
    """Safely load data from MongoDB collection"""
    try:
        collection = mydb[collection_name]
        # Exclude MongoDB's _id field from results
        data = list(collection.find())
        streamlit_logger.debug(f"Successfully loaded {len(data)} records from {collection_name}")
        return data
    except Exception as e:
        streamlit_logger.error(f"Error loading data from MongoDB: {str(e)}")
        return []

def save_data(data: List[Dict], collection_name: str) -> None:
    """Safely save data to MongoDB collection"""
    try:
        collection = mydb[collection_name]
        
        result = collection.insert_one(data[-1])  # Use copy to avoid any modification issues
        streamlit_logger.info(f"Insert result: {result.inserted_id}")
            
        # Verify insertion
        if result.inserted_id:
            found = collection.find_one({"_id": result.inserted_id})
            if found:
                streamlit_logger.info("Document verified in database")
                return True
            else:
                streamlit_logger.error("Document not found after insertion")
                return False
        return False
        
    except Exception as e:
        streamlit_logger.error(f"Save error: {str(e)}", exc_info=True)
        return False

def serialize_pydantic_model(model) -> Dict:
    """Convert Pydantic model to serializable dict"""
    try:
        if isinstance(model, BaseModel):
            return model.model_dump()
        return model
    except Exception as e:
        streamlit_logger.error(f"Error serializing model: {str(e)}")
        raise ValueError(f"Could not serialize model: {str(e)}")


def save_parsed_resume(structured_data: dict, resume_text: str, filename: str) -> None:
    """Save parsed resume data to MongoDB"""
    try:
        global resume_jsn_data
        # Convert all Pydantic models to serializable dicts
        serializable_data = {
            section_name: serialize_pydantic_model(section_content)
            for section_name, section_content in structured_data.items()
        }
        
        # Create the complete record
        resume_record = {
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "parsed_data": serializable_data,
            "raw_text_sample": resume_text[:1000] + ("..." if len(resume_text) > 1000 else "")
        }
        # Save directly to MongoDB (no list handling needed)
        save_data([resume_record], "job_application")
        
    except Exception as e:
        streamlit_logger.error(f"Failed to save resume data: {str(e)}")
        raise

def extract_text_from_pdf(file_path_or_object):
    try:
        if isinstance(file_path_or_object, str):
            # It's a file path
            with open(file_path_or_object, 'rb') as file:
                reader = PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
        else:
            # It's a file object
            reader = PdfReader(file_path_or_object)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return text.strip()
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_path_or_object):
    try:
        if isinstance(file_path_or_object, str):
            # It's a file path
            doc = Document(file_path_or_object)
        else:
            # It's a file object
            doc = Document(file_path_or_object)
            
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        return "\n\n".join(paragraphs)
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {str(e)}")
        return ""