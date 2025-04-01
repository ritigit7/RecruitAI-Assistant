# ======================================================================================================
#                                               imports        
# ======================================================================================================

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from openai import OpenAI
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from pydantic import BaseModel, Field
from PyPDF2 import PdfReader
from docx import Document
import json
import datetime
import logging
import os
from typing import Optional, List, Dict
from datetime import datetime, timedelta
import pandas as pd
import calendar
import pymongo
from pdf_maker import pdfmaker
import tempfile
import base64

# ======================================================================================================
#                                       logging configuration       
# ======================================================================================================

# Create logs directory if it doesn't exist
os.makedirs("logs", exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/app.log"),
        logging.StreamHandler()
    ]
)

# Create specialized loggers
ollama_logger = logging.getLogger('ollama')
resume_logger = logging.getLogger('resume')
meeting_logger = logging.getLogger('meeting')
flask_logger = logging.getLogger('flask')

# ======================================================================================================
#                                       ollama llm connect        
# ======================================================================================================

ollama_client = OpenAI(
    base_url="http://localhost:11434/v1",
    api_key="ollama"
)

ollama_logger.info("Initializing Ollama client with base URL: http://localhost:11434/v1")

model = "llama3.2"
ollama_logger.info(f"Using model: {model}")

message = [
    {
        "role": "system",
        "content": """You are a highly accurate and comprehensive resume builder AI assistant. Your ideal tool should be able to analyze resume and 
provide personalized feedback on the following areas:

1.  Personal Details: Suggest Full Name, contact information, Mail ,github or linked profile.
2.  Professional Summary: Offer a rewritten summary that showcases About Me , Summary.
3.  Work Experience: identify gaps in employment history, suggest rewording or reorganizing of job descriptions, and provide insights 
    on how to highlight key accomplishments.
4.  Education: Verify the accuracy of educational credentials, suggest degree information, and recommend relevant coursework 
    or specializations.
5.  Skills: Identify all skills in list of skills, suggest adding new skills that align with profession or industry, and provide 
    guidance on how to showcase them effectively.
6.  Projects: find projects or experiences in resume include , along with a brief description of each project's 
    significance.
7.  Hobbies: get hobbies or interests that can demonstrate personality, skills, or character traits to potential employers.
8.  Achievements: identify any notable achievements, awards, or publications."""
    },
]

ollama_logger.info("Setting up system message for resume analysis")

try:
    response = ollama_client.chat.completions.create(
        model=model,
        messages=message,
    )
    ollama_logger.info("Initial Ollama connection successful")
except Exception as e:
    ollama_logger.error(f"Failed to connect to Ollama: {str(e)}")

# ======================================================================================================
#                                   Pydantic response format for RESUME
# ======================================================================================================

resume_logger.info("Defining Pydantic models for resume parsing")

class ApplicationForm(BaseModel):
    date: str = Field(description="Date of application")
    site: str = Field(description="Site of application")
    GET_applied_for: str = Field(description="GET applied for")

class NameDetails(BaseModel):
    first_name: str
    middle_name: str
    last_name: str

class ResidenceDuration(BaseModel):
    years: int
    months: int

class VehicleDetails(BaseModel):
    has_vehicle: bool
    vehicle_type: Optional[str] = None

class PersonalInformation(BaseModel):
    name: NameDetails
    date_of_birth: str
    age: int
    gender: str
    blood_group: str
    marital_status: str
    marriage_date: Optional[str] = None
    religion: str
    caste: str
    contact_number: str
    alternate_contact_number: Optional[str] = None
    personal_email: str
    height: float
    weight: float
    aadhar_no: str
    pan_no: str
    residence_status: str
    present_address: str
    permanent_address: str
    residence_duration: ResidenceDuration
    distance_from_office_km: float
    own_vehicle: VehicleDetails

class FamilyMember(BaseModel):
    relationship: str
    name: str
    occupation: str
    age: int
    dependant: bool

class FamilyDetails(BaseModel):
    family_members_count: int
    members: List[FamilyMember]

class EducationDetails(BaseModel):
    qualification: str
    college_school: str
    board_university: str
    percentage_grade: str
    year_of_passing: str

class LanguageProficiency(BaseModel):
    speak: bool
    read: bool
    write: bool

class LanguagesKnown(BaseModel):
    english: LanguageProficiency
    gujarati: LanguageProficiency
    hindi: LanguageProficiency

class WorkExperience(BaseModel):
    employer_name: str
    location: str
    designation: str
    start_date: str
    end_date: str
    salary: str
    reason_for_leaving: str

class ExperienceDetails(BaseModel):
    current_salary_ctc: str
    expected_salary_ctc: str
    notice_period_months: int
    work_experience: List[WorkExperience]

class ReferenceDetails(BaseModel):
    name: str
    company_name: str
    occupation: str
    contact_number: str

class ITSkillsCertifications(BaseModel):
    skill_course: str
    institute: str
    year: str

class GeneralDetails(BaseModel):
    hobbies: str
    ready_to_sign_bond: bool
    it_skills_certifications: List[ITSkillsCertifications]
    other_information: Optional[str] = None
    training_internship_details: Optional[str] = None
    professional_memberships: Optional[str] = None

class CandidateSignature(BaseModel):
    full_name: str
    date: str
    place: str
    signature: str

class TechnicalInterview(BaseModel):
    interviewer_name: str
    start_time: str
    end_time: str
    date: str


# ======================================================================================================
#                                   Pydantic response format for Meeting
# ======================================================================================================

meeting_logger.info("Defining Pydantic models for meeting extraction")

class MeetingInfo(BaseModel):
    """Extracted information about whether the user input relates to a meeting."""
    meeting_description: str = Field(description="Brief description of the meeting purpose.")
    is_calendar_event: bool = Field(description="Indicates if the input is a calendar event.")
    confidence_score: float = Field(description="Confidence level (0 to 1) for meeting detection.")
    title: str = Field(description="The title or subject of the meeting.")
    datetime: str = Field(description="The date and time of the meeting.format example: 2025-03-26T14:00:00")
    date: str = Field(description="The date of the meeting in match format %A, %yyyy-%mm-%dd.")
    time: str = Field(description="The time of the meeting in match format  %HH:%MM:%SS.")
    duration_of_meeting: float = Field(description="Duration of the meeting in hours.")
    participants: List[str] = Field(description="List of participants attending the meeting.")
    confirmation_message: str = Field(description="Natural language confirmation message for the user.")
    calendar_link: Optional[str] = Field(description="Generated calendar link, if applicable.")

# ======================================================================================================
#                                     similarity search prompts        
# ======================================================================================================

resume_logger.info("Setting up similarity search prompts for resume sections")

msg1 = "Extract application form details, including the date, site, and the GET applied for. These details are typically found at the beginning of the application form."
msg2 = "Extract name details, including the first name, middle name, and last name from the application form."
msg3 = "Extract residence duration details, including the number of years and months the candidate has lived at their current residence."
msg4 = "Extract vehicle ownership details, including whether the candidate owns a vehicle and the type of vehicle they have."
msg5 = "Extract personal information, including date of birth, age, gender, blood group, marital status, marriage date, religion, caste, contact numbers, email, height, weight, Aadhar number, PAN number, residence status, present and permanent address, and distance from the office."
msg6 = "Extract individual family member details, including the relationship (father, mother, spouse, sibling/child), name, occupation, age, and whether they are a dependent."
msg7 = "Extract overall family details, including the total number of family members and a list of individual family members with their details."
msg8 = "Extract education details, including qualification, college/school name, board/university, percentage/grade, and year of passing. Higher education details should be prioritized."
msg9 = "Extract language proficiency details, including whether the candidate can speak, read, or write in different languages."
msg10 = "Extract details about languages known by the candidate, including proficiency levels for English, Hindi, and Gujarati."
msg11 = "Extract work experience details, including employer names, locations, job designations, start and end dates, salary details, and reasons for leaving previous jobs."
msg12 = "Extract overall experience details, including current salary CTC, expected salary CTC, notice period, and a structured list of previous work experiences."
msg13 = "Extract reference details, including the name, company name, occupation, and contact number of professional references provided by the candidate."
msg14 = "Extract IT skills and certifications, including skill/course name, the institute where it was obtained, and the year of completion."
msg15 = "Extract general details, including hobbies, willingness to sign a bond, IT skills, other relevant information, training or internship details, and professional memberships."
msg16 = "Extract candidate signature details, including full name, date, place, and digital or physical signature."
msg17 = "Extract technical interview details, including the interviewer's name, start time, end time, and interview date."

msg = [msg1, msg2, msg3, msg4, msg5, msg6, msg7, msg8, msg9, msg10, msg11, msg12, msg13, msg14, msg15, msg16, msg17]

# ======================================================================================================
#                             get llm response using system prompt and classes for RESUME     
# ======================================================================================================

def get_llm_response(userInput: str, message:str, response_format: type[BaseModel]) -> BaseModel:
    try:
        resume_logger.info(f"Extracting {response_format.__name__} details")
        response = ollama_client.beta.chat.completions.parse(
            model=model,
            messages=[
                {"role": "system", "content": message},
                {"role": "user", "content": userInput} 
            ],
            response_format=response_format
        )
        resume_logger.info(f"Successfully extracted {response_format.__name__} details")
        return response.choices[0].message.parsed
    except Exception as e:
        resume_logger.error(f"Error getting LLM response: {str(e)}")
        raise

def fn_ApplicationForm(userInput: str) -> ApplicationForm:
    return get_llm_response(userInput, msg1, ApplicationForm)

def fn_NameDetails(userInput: str) -> NameDetails:
    return get_llm_response(userInput, msg2, NameDetails)

def fn_ResidenceDuration(userInput: str) -> ResidenceDuration:
    return get_llm_response(userInput, msg3, ResidenceDuration)

def fn_VehicleDetails(userInput: str) -> VehicleDetails:
    return get_llm_response(userInput, msg4, VehicleDetails)

def fn_PersonalInformation(userInput: str) -> PersonalInformation:
    return get_llm_response(userInput, msg5, PersonalInformation)

def fn_FamilyMember(userInput: str) -> FamilyMember:
    return get_llm_response(userInput, msg6, FamilyMember)

def fn_FamilyDetails(userInput: str) -> FamilyDetails:
    return get_llm_response(userInput, msg7, FamilyDetails)

def fn_EducationDetails(userInput: str) -> EducationDetails:
    return get_llm_response(userInput, msg8, EducationDetails)

def fn_LanguageProficiency(userInput: str) -> LanguageProficiency:
    return get_llm_response(userInput, msg9, LanguageProficiency)

def fn_LanguagesKnown(userInput: str) -> LanguagesKnown:
    return get_llm_response(userInput, msg10, LanguagesKnown)

def fn_WorkExperience(userInput: str) -> WorkExperience:
    return get_llm_response(userInput, msg11, WorkExperience)

def fn_ExperienceDetails(userInput: str) -> ExperienceDetails:
    return get_llm_response(userInput, msg12, ExperienceDetails)

def fn_ReferenceDetails(userInput: str) -> ReferenceDetails:
    return get_llm_response(userInput, msg13, ReferenceDetails)

def fn_ITSkillsCertifications(userInput: str) -> ITSkillsCertifications:
    return get_llm_response(userInput, msg14, ITSkillsCertifications)

def fn_GeneralDetails(userInput: str) -> GeneralDetails:
    return get_llm_response(userInput, msg15, GeneralDetails)

def fn_CandidateSignature(userInput: str) -> CandidateSignature:
    return get_llm_response(userInput, msg16, CandidateSignature)

def fn_TechnicalInterview(userInput: str) -> TechnicalInterview:
    return get_llm_response(userInput, msg17, TechnicalInterview)

def fn_Resume(resume_text: str) -> list:
    resume_logger.info("Starting full resume analysis")
    
    if not resume_text.strip():
        resume_logger.warning("Empty resume text provided")
        return []

    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=500,
            separators=["\n\n", "\n", " ", ""]
        )

        documents = text_splitter.split_text(resume_text)
        resume_logger.info(f"Split resume into {len(documents)} chunks")

        # Use a dummy embedding if no text is available
        if not documents:
            resume_logger.warning("No document chunks were created, using empty text for embeddings")
            db = FAISS.from_texts([""], OllamaEmbeddings(model="nomic-embed-text"))
        else:
            resume_logger.info("Creating FAISS vector store with nomic-embed-text embeddings")
            db = FAISS.from_texts(documents, OllamaEmbeddings(model="nomic-embed-text"))

        results = []
        for i, query in enumerate(msg):
            resume_logger.info(f"Performing similarity search for section {i+1}")
            rs = db.similarity_search(query, k=5)
            tx = " ".join([doc.page_content for doc in rs])
            results.append(tx)
            resume_logger.debug(f"Section {i+1} search results length: {len(tx)} characters")
        
        # Extract all sections
        sections = {
            "Application Form": fn_ApplicationForm(results[0]),
            "Name Details": fn_NameDetails(results[1]),
            "Residence Duration": fn_ResidenceDuration(results[2]),
            "Vehicle Details": fn_VehicleDetails(results[3]),
            "Personal Information": fn_PersonalInformation(results[4]),
            "Family Member": fn_FamilyMember(results[5]),
            "Family Details": fn_FamilyDetails(results[6]),
            "Education Details": fn_EducationDetails(results[7]),
            "Language Proficiency": fn_LanguageProficiency(results[8]),
            "Languages Known": fn_LanguagesKnown(results[9]),
            "Work Experience": fn_WorkExperience(results[10]),
            "Experience Details": fn_ExperienceDetails(results[11]),
            "Reference Details": fn_ReferenceDetails(results[12]),
            "IT Skills & Certifications": fn_ITSkillsCertifications(results[13]),
            "General Details": fn_GeneralDetails(results[14]),
            "Candidate Signature": fn_CandidateSignature(results[15]),
            "Technical Interview": fn_TechnicalInterview(results[16]),
        }
        
        return sections

    except Exception as e:
        resume_logger.error(f"Error in complete resume analysis: {str(e)}")
        raise

# ======================================================================================================
#                          get llm response using system prompt and classes for Meeting
# ======================================================================================================

def extract_meeting_info(user_input: str) -> Optional[MeetingInfo]:
    """Extracts structured meeting information using FAISS and LLM."""
    meeting_logger.info("Starting meeting extraction process")
    
    try:
        meeting_logger.info("Calling Ollama model for meeting extraction")
        response = ollama_client.beta.chat.completions.parse(
            model=model,
            messages=[
                {
                    'role': 'system',
                    'content': '''
                        You are a meeting information extractor. Your task is to find meeting details in user text and put them into a specific JSON format.
                        Follow these instructions strictly and return only the JSON.

                        Expected JSON Format:
                        {
                            "meeting_description": "Brief description of the meeting purpose.",
                            "is_calendar_event": true/false,
                            "confidence_score": float (0 to 1),
                            "title": "The title or subject of the meeting.",
                            "datetime": "YYYY-MM-DDTHH:MM:SS",
                            "date": "Day, YYYY-MM-DD",
                            "time": "HH:MM:SS",
                            "duration_of_meeting": float (in hours),
                            "participants": ["List", "of", "participants"],
                            "confirmation_message": "Natural language confirmation message for the user.",
                            "calendar_link": "Generated calendar link if applicable, else null."
                        }

                        Ensure accuracy in parsing and extracting details. If any field is missing or unclear, make a reasonable assumption.
                    '''
                },
                {'role': 'user', 'content': user_input}
            ],
            response_format=MeetingInfo
        )

        info = response.choices[0].message.parsed
        meeting_logger.info(f"LLM extraction successful - Title: '{info.title}', Date: '{info.date}', Time: '{info.time}'")

        if not info or not info.is_calendar_event or info.confidence_score < 0.6:
            meeting_logger.warning("No valid meeting detected or confidence too low")
            return None
        
        meeting_logger.info("Meeting extraction completed successfully")
        # Final aggregated data
        return {
            "meeting_description": info.meeting_description,
            "is_calendar_event": info.is_calendar_event,
            "confidence_score": info.confidence_score,
            "title": info.title,
            "datetime": info.datetime,
            "date": info.date,
            "time": info.time,
            "duration_of_meeting": info.duration_of_meeting,
            "participants": info.participants,
            "confirmation_message": info.confirmation_message,
            "calendar_link": info.calendar_link
        }
    except Exception as e:
        meeting_logger.error(f"Error in meeting extraction process: {str(e)}")
        raise

# ======================================================================================================
#                                       DATA STORAGE CONFIGURATION in MongoDB
# ======================================================================================================

MEETINGS_FILE = "meeting_records"
RESUMES_FILE = "job_application"

myclient = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = myclient["resumedb"]

def load_data(collection_name: str) -> List[Dict]:
    """Safely load data from MongoDB collection"""
    try:
        collection = mydb[collection_name]
        # Exclude MongoDB's _id field from results
        data = list(collection.find())
        flask_logger.debug(f"Successfully loaded {len(data)} records from {collection_name}")
        return data
    except Exception as e:
        flask_logger.error(f"Error loading data from MongoDB: {str(e)}")
        return []

def save_data(data: List[Dict], collection_name: str) -> None:
    """Safely save data to MongoDB collection"""
    try:
        collection = mydb[collection_name]
        
        result = collection.insert_one(data[-1])  # Use copy to avoid any modification issues
        flask_logger.info(f"Insert result: {result.inserted_id}")
            
        # Verify insertion
        if result.inserted_id:
            found = collection.find_one({"_id": result.inserted_id})
            if found:
                flask_logger.info("Document verified in database")
                return True
            else:
                flask_logger.error("Document not found after insertion")
                return False
        return False
        
    except Exception as e:
        flask_logger.error(f"Save error: {str(e)}", exc_info=True)
        return False

def serialize_pydantic_model(model) -> Dict:
    """Convert Pydantic model to serializable dict"""
    try:
        if isinstance(model, BaseModel):
            return model.model_dump()
        return model
    except Exception as e:
        flask_logger.error(f"Error serializing model: {str(e)}")
        raise ValueError(f"Could not serialize model: {str(e)}")

# ======================================================================================================
#                                       File Text Extraction Functions
# ======================================================================================================

def save_parsed_resume(structured_data: dict, resume_text: str, filename: str) -> None:
    """Save parsed resume data to MongoDB"""
    try:
        serializable_data = {
            section_name: serialize_pydantic_model(section_content)
            for section_name, section_content in structured_data.items()
        }
        
        # Create the complete record
        resume_record = {
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "parsed_data": serializable_data,
            "raw_text_sample": resume_text[:1000] + ("..." if len(resume_text) > 1000 else "")
        }
        # Save directly to MongoDB (no list handling needed)
        save_data([resume_record], "job_application")

    except Exception as e:
        raise ValueError(f"Error saving parsed resume data: {str(e)}")


def extract_text_from_pdf(file_path):
    try:
        flask_logger.info(f"Extracting text from PDF file: {file_path}")
        reader = PdfReader(file_path)
        text = ""
        for page_num, page in enumerate(reader.pages):
            flask_logger.debug(f"Processing page {page_num+1}")
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        flask_logger.info(f"Successfully extracted text from PDF file: {file_path}")
        return text.strip()
    except Exception as e:
        flask_logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file_path):
    try:
        flask_logger.info(f"Extracting text from DOCX file: {file_path}")
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        table_count = len(doc.tables)
        flask_logger.debug(f"Found {table_count} tables in document")
        for table_idx, table in enumerate(doc.tables):
            flask_logger.debug(f"Processing table {table_idx+1}")
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        flask_logger.info(f"Successfully extracted text from DOCX file: {file_path}")
        return "\n\n".join(paragraphs)
    except Exception as e:
        flask_logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

# ======================================================================================================
#                                       Flask Application Setup
# ======================================================================================================

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes

@app.route('/parse-resume', methods=['GET'])
def parse_resume_endpoint():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    # Save the file temporarily
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, file.filename)
    file.save(file_path)
    
    try:
        # Extract text based on file type
        if file.filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file.filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            return jsonify({'error': 'Unsupported file type'}), 400
        
        # Parse the resume using your existing function
        parsed_data = fn_Resume(text)
        
        # Convert Pydantic models to dicts for JSON serialization
        serializable_data = {
            section_name: serialize_pydantic_model(section_content)
            for section_name, section_content in parsed_data.items()
        }
        
        return jsonify(serializable_data)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up temporary files
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(temp_dir):
            os.rmdir(temp_dir)

@app.route('/extract-meeting', methods=['GET'])
def extract_meeting_endpoint():
    data = request.get_json()
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    
    try:
        # Use your existing meeting extraction function
        meeting_info = extract_meeting_info(data['text'])
        if meeting_info is None:
            return jsonify({'error': 'No valid meeting detected'}), 400
        return jsonify(meeting_info)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/generate-pdf', methods=['GET'])
def generate_pdf_endpoint():
    data = request.get_json()
    if not data or 'resume_data' not in data:
        return jsonify({'error': 'No resume data provided'}), 400
    
    try:
        # Use your existing PDF generation function
        pdf_path = pdfmaker(data['resume_data'])
        
        # Return the PDF file
        return send_file(
            pdf_path,
            as_attachment=True,
            download_name='resume.pdf',
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)