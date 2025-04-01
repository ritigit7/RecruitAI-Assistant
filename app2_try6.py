# ======================================================================================================
#                                               imports        
# ======================================================================================================

from openai import OpenAI
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from pydantic import BaseModel, Field
import streamlit as st
from PyPDF2 import PdfFileReader
from docx import Document
import json
import datetime
import logging
import os
from typing import Optional, List, Dict
from datetime import datetime, timedelta
import pandas as pd
import plotly.express as px
import calendar
import pymongo
from pdf_maker import pdfmaker
import base64


# ======================================================================================================
#                                       logging configuration       
# ======================================================================================================

# Create logs directory if it doesn't exist
os.makedirs("logs", exist_ok=True)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        # logging.FileHandler("logs/app.log"),
        logging.StreamHandler()
    ]
)

# Create specialized loggers
ollama_logger = logging.getLogger('ollama')
resume_logger = logging.getLogger('resume')
meeting_logger = logging.getLogger('meeting')
streamlit_logger = logging.getLogger('streamlit')

# ======================================================================================================
#                                       ollama llm connect        
# ======================================================================================================

ollama_client = OpenAI(
    base_url="http://localhost:11434/v1",
    api_key="ollama"
)

ollama_logger.info("Initializing Ollama client with base URL: http://localhost:11434/v1")

# model = "llama3.2"
model = "llama3.1"
# model = "mistral"
ollama_logger.info(f"Using model: {model}")

message = [
    {
        "role": "system",
        "content": """You are a highly accurate and comprehensive resume builder AI assistant. Your ideal tool should be able to analyze resume and 
provide personalized feedback on the following areas:

1.  Personal Details: Suggest Full Name, contact information, Mail ,github or linked profile.
2.  Professional Summary: Offer a rewritten summary that showcases About Me , Summary.
3.  Work Experience: identify gaps in employment history, suggest rewording or reorganizing of job descriptions, and provide insights 
    on how to highlight key accomplishments.
4.  Education: Verify the accuracy of educational credentials, suggest degree information, and recommend relevant coursework 
    or specializations.
5.  Skills: Identify all skills in list of skills, suggest adding new skills that align with profession or industry, and provide 
    guidance on how to showcase them effectively.
6.  Projects: find projects or experiences in resume include , along with a brief description of each project's 
    significance.
7.  Hobbies: get hobbies or interests that can demonstrate personality, skills, or character traits to potential employers.
8.  Achievements: identify any notable achievements, awards, or publications."""
    },
]

ollama_logger.info("Setting up system message for resume analysis")

try:
    response = ollama_client.chat.completions.create(
        model=model,
        messages=message,
    )
    ollama_logger.info("Initial Ollama connection successful")
except Exception as e:
    ollama_logger.error(f"Failed to connect to Ollama: {str(e)}")
    st.error(f"Failed to connect to Ollama: {str(e)}")

# ======================================================================================================
#                                   Pydantic response format for RESUME
# ======================================================================================================

resume_logger.info("Defining Pydantic models for resume parsing")

class ApplicationForm(BaseModel):
    date: str = Field(description="Date of application")
    site: str = Field(description="Site of application")
    post_applied_for: str = Field(description="Post applied for")

class NameDetails(BaseModel):
    first_name: str
    middle_name: str
    last_name: str

class ResidenceDuration(BaseModel):
    years: int
    months: int

class VehicleDetails(BaseModel):
    has_vehicle: bool
    vehicle_type: Optional[str] = None

class PersonalInformation(BaseModel):
    name: NameDetails
    date_of_birth: str
    age: int
    gender: str
    blood_group: str
    marital_status: str
    marriage_date: Optional[str] = None
    religion: str
    caste: str
    contact_number: str
    alternate_contact_number: Optional[str] = None
    personal_email: str
    height: float
    weight: float
    aadhar_no: str
    pan_no: str
    residence_status: str
    present_address: str
    permanent_address: str
    residence_duration: ResidenceDuration
    distance_from_office_km: float
    own_vehicle: VehicleDetails

class FamilyMember(BaseModel):
    relationship: str
    name: str
    occupation: str
    age: int
    dependant: bool

class FamilyDetails(BaseModel):
    family_members_count: int
    members: List[FamilyMember]

class EducationDetails(BaseModel):
    qualification: str
    college_school: str
    board_university: str
    percentage_grade: str
    year_of_passing: str

class LanguageProficiency(BaseModel):
    speak: bool
    read: bool
    write: bool

class LanguagesKnown(BaseModel):
    english: LanguageProficiency
    gujarati: LanguageProficiency
    hindi: LanguageProficiency

class WorkExperience(BaseModel):
    employer_name: str
    location: str
    designation: str
    start_date: str
    end_date: str
    salary: str
    reason_for_leaving: str

class ExperienceDetails(BaseModel):
    current_salary_ctc: str
    expected_salary_ctc: str
    notice_period_months: int
    work_experience: List[WorkExperience]

class ReferenceDetails(BaseModel):
    name: str
    company_name: str
    occupation: str
    contact_number: str

class ITSkillsCertifications(BaseModel):
    skill_course: str
    institute: str
    year: str

class GeneralDetails(BaseModel):
    hobbies: str
    ready_to_sign_bond: bool
    it_skills_certifications: List[ITSkillsCertifications]
    other_information: Optional[str] = None
    training_internship_details: Optional[str] = None
    professional_memberships: Optional[str] = None

class OfficeUseOnly(BaseModel):
    technical_comments: str
    technical_interviewer_name: str
    interview_start_time:  str
    interview_end_time: str
    date_of_interview: str

# ======================================================================================================
#                                   Pydantic response format for Meeting
# ======================================================================================================

meeting_logger.info("Defining Pydantic models for meeting extraction")

class MeetingInfo(BaseModel):
    """Extracted information about whether the user input relates to a meeting."""
    meeting_description: str = Field(description="Brief description of the meeting purpose.")
    is_calendar_event: bool = Field(description="Indicates if the input is a calendar event.")
    confidence_score: float = Field(description="Confidence level (0 to 1) for meeting detection.")
    title: str = Field(description="The title or subject of the meeting.")
    datetime: str = Field(description="The date and time of the meeting.format example: 2025-03-26T14:00:00")
    date: str = Field(description="The date of the meeting in match format %A, %yyyy-%mm-%dd.")
    time: str = Field(description="The time of the meeting in match format  %HH:%MM:%SS.")
    duration_of_meeting: float = Field(description="Duration of the meeting in hours.")
    participants: List[str] = Field(description="List of participants attending the meeting.")
    confirmation_message: str = Field(description="Natural language confirmation message for the user.")
    calendar_link: Optional[str] = Field(description="Generated calendar link, if applicable.")
    notes: Optional[str] = Field(description="Additional notes or details about the meeting.")
    location: Optional[str] = Field(description="Location of the meeting.")

# ======================================================================================================
#                                     similarity search prompts        
# ======================================================================================================

resume_logger.info("Setting up similarity search prompts for resume sections")

msg1 = "Extract application form details, including the date, site, and the post applied for. These details are typically found at the beginning of the application form. if you don't find the answer remain empty the field."
msg2 = "Extract name details, including the first name, middle name, and last name from the application form. if you don't find the answer remain empty the field."
msg3 = "Extract residence duration details, including the number of years and months the candidate has lived at their current residence. if you don't find the answer remain empty the field."
msg4 = "Extract vehicle ownership details, including whether the candidate owns a vehicle and the type of vehicle they have. if you don't find the answer remain empty the field."
msg5 = "Extract personal information, including date of birth, age, gender, blood group, marital status, marriage date, religion, caste, contact numbers, email, height, weight, Aadhar number, PAN number, residence status, present and permanent address, and distance from the office. if you don't find the answer remain empty the field."
msg6 = "Extract individual family member details, including the relationship (father, mother, spouse, sibling/child), name, occupation, age, and whether they are a dependent. if you don't find the answer remain empty the field."
msg7 = "Extract overall family details, including the total number of family members and a list of individual family members with their details. if you don't find the answer remain empty the field."
msg8 = "Extract education details, including qualification, college/school name, board/university, percentage/grade, and year of passing. Higher education details should be prioritized. if you don't find the answer remain empty the field."
msg9 = "Extract language proficiency details, including whether the candidate can speak, read, or write in different languages. if you don't find the answer remain empty the field."
msg10 = "Extract details about languages known by the candidate, including proficiency levels for English, Hindi, and Gujarati. if you don't find the answer remain empty the field."
msg11 = "Extract work experience details, including employer names, locations, job designations, start and end dates, salary details, and reasons for leaving previous jobs. if you don't find the answer remain empty the field."
msg12 = "Extract overall experience details, including current salary CTC, expected salary CTC, notice period, and a structured list of previous work experiences. if you don't find the answer remain empty the field."
msg13 = "Extract reference details, including the name, company name, occupation, and contact number of professional references provided by the candidate. if you don't find the answer remain empty the field."
msg14 = "Extract IT skills and certifications, including skill/course name, the institute where it was obtained, and the year of completion. if you don't find the answer remain empty the field."
msg15 = "Extract general details, including hobbies, willingness to sign a bond, IT skills, other relevant information, training or internship details, and professional memberships. if you don't find the answer remain empty the field."
msg16 = "Extract candidate signature details, including full name, date, place, and digital or physical signature. if you don't find the answer remain empty the field."
msg17 = "Extract technical interview details, including the interviewer's name, start time, end time, and interview date. if you don't find the answer remain empty the field."
msg18 = "Extract all application data, including personal information, family details, education, work experience, references, certifications, general details, and technical interview information. if you don't find the answer remain empty the field."

msg = [msg1, msg2, msg3, msg4, msg5, msg6, msg7, msg8, msg9, msg10, msg11, msg12, msg13, msg14, msg15, msg16, msg17, msg18]

# ======================================================================================================
#                             get llm response using system prompt and classes for RESUME     
# ======================================================================================================


def get_llm_response(userInput: str, message:str,response_format: type[BaseModel]) -> BaseModel:
    with st.spinner(f"Extracting details of {response_format.__name__}"):
        try:
            resume_logger.info(f"Extracting {response_format.__name__} details")
            response = ollama_client.beta.chat.completions.parse(
                model=model,
                messages=[
                    {"role": "system", "content": message},
                    {"role": "user", "content": userInput} 
                ],
                response_format=response_format
            )
            resume_logger.info(f"Successfully extracted {response_format.__name__} details")
            return response.choices[0].message.parsed
        except Exception as e:
            resume_logger.error(f"Error getting LLM response: {str(e)}")
            raise

def fn_ApplicationForm(userInput: str) -> ApplicationForm:
    return get_llm_response(userInput,msg1, ApplicationForm)

def fn_NameDetails(userInput: str) -> NameDetails:
    return get_llm_response(userInput,msg2, NameDetails)

def fn_ResidenceDuration(userInput: str) -> ResidenceDuration:
    return get_llm_response(userInput,msg3, ResidenceDuration)

def fn_VehicleDetails(userInput: str) -> VehicleDetails:
    return get_llm_response(userInput,msg4, VehicleDetails)

def fn_PersonalInformation(userInput: str) -> PersonalInformation:
    return get_llm_response(userInput,msg5, PersonalInformation)

def fn_FamilyMember(userInput: str) -> FamilyMember:
    return get_llm_response(userInput,msg6, FamilyMember)

def fn_FamilyDetails(userInput: str) -> FamilyDetails:
    return get_llm_response(userInput,msg7, FamilyDetails)

def fn_EducationDetails(userInput: str) -> EducationDetails:
    return get_llm_response(userInput,msg8, EducationDetails)

def fn_LanguageProficiency(userInput: str) -> LanguageProficiency:
    return get_llm_response(userInput,msg9, LanguageProficiency)

def fn_LanguagesKnown(userInput: str) -> LanguagesKnown:
    return get_llm_response(userInput,msg10, LanguagesKnown)

def fn_WorkExperience(userInput: str) -> WorkExperience:
    return get_llm_response(userInput,msg11, WorkExperience)

def fn_ExperienceDetails(userInput: str) -> ExperienceDetails:
    return get_llm_response(userInput,msg12, ExperienceDetails)

def fn_ReferenceDetails(userInput: str) -> ReferenceDetails:
    return get_llm_response(userInput,msg13, ReferenceDetails)

def fn_ITSkillsCertifications(userInput: str) -> ITSkillsCertifications:
    return get_llm_response(userInput,msg14, ITSkillsCertifications)

def fn_GeneralDetails(userInput: str) -> GeneralDetails:
    return get_llm_response(userInput,msg15, GeneralDetails)

def fn_OfficeUseOnly(userInput: str) -> OfficeUseOnly:
    return get_llm_response(userInput,msg16, OfficeUseOnly)


def fn_Resume(resume_text: str) -> list:
    resume_logger.info("Starting full resume analysis")
    
    if not resume_text.strip():
        resume_logger.warning("Empty resume text provided")
        return []

    try:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1500,
            chunk_overlap=300,
            separators=["\n\n", "\n", " ", ""]
        )

        documents = text_splitter.split_text(resume_text)
        resume_logger.info(f"Split resume into {len(documents)} chunks")

        # Use a dummy embedding if no text is available
        if not documents:
            resume_logger.warning("No document chunks were created, using empty text for embeddings")
            db = FAISS.from_texts([""], OllamaEmbeddings(model="nomic-embed-text"))
        else:
            resume_logger.info("Creating FAISS vector store with nomic-embed-text embeddings")
            db = FAISS.from_texts(documents, OllamaEmbeddings(model="nomic-embed-text"))

        results = []
        for i, query in enumerate(msg):
            resume_logger.info(f"Performing similarity search for section {i+1}")
            rs = db.similarity_search(query, k=7)
            tx = " ".join([doc.page_content for doc in rs])
            results.append(tx)
            resume_logger.debug(f"Section {i+1} search results length: {len(tx)} characters")
        
        # Extract all sections
        sections = {
            "Application Form": fn_ApplicationForm(results[0]),
            "Name Details": fn_NameDetails(results[1]),
            "Residence Duration": fn_ResidenceDuration(results[2]),
            "Vehicle Details": fn_VehicleDetails(results[3]),
            "Personal Information": fn_PersonalInformation(results[4]),
            "Family Member": fn_FamilyMember(results[5]),
            "Family Details": fn_FamilyDetails(results[6]),
            "Education Details": fn_EducationDetails(results[7]),
            "Language Proficiency": fn_LanguageProficiency(results[8]),
            "Languages Known": fn_LanguagesKnown(results[9]),
            "Work Experience": fn_WorkExperience(results[10]),
            "Experience Details": fn_ExperienceDetails(results[11]),
            "Reference Details": fn_ReferenceDetails(results[12]),
            "IT Skills & Certifications": fn_ITSkillsCertifications(results[13]),
            "General Details": fn_GeneralDetails(results[14]),
            "Office Use Only": fn_OfficeUseOnly(results[15])
        }
        
        return sections

    except Exception as e:
        resume_logger.error(f"Error in complete resume analysis: {str(e)}")
        raise

# ======================================================================================================
#                          get llm response using system prompt and classes for Meeting
# ======================================================================================================

def extract_meeting_info(user_input: str) -> Optional[MeetingInfo]:
    """Extracts structured meeting information using FAISS and LLM."""
    meeting_logger.info("Starting meeting extraction process")
    
    try:
        meeting_logger.info("Calling Ollama model for meeting extraction")
        response = ollama_client.beta.chat.completions.parse(
            model=model,
            messages=[
                {
                    'role': 'system',
                    'content': '''
                        You are a meeting information extractor. Your task is to find meeting details in user text and put them into a specific JSON format.
                        Follow these instructions strictly and return only the JSON.

                        Expected JSON Format:
                        {
                            "meeting_description": "Brief description of the meeting purpose.",
                            "is_calendar_event": true/false,
                            "confidence_score": float (0 to 1),
                            "title": "The title or subject of the meeting.",
                            "datetime": "YYYY-MM-DDTHH:MM:SS",
                            "date": "Day, YYYY-MM-DD",
                            "time": "HH:MM:SS",
                            "duration_of_meeting": float (in hours),
                            "participants": ["List", "of", "participants"],
                            "confirmation_message": "Natural language confirmation message for the user.",
                            "calendar_link": "Generated calendar link if applicable, else null."
                            "notes": "Additional notes or details about the meeting."
                            "location":"Location of the meeting. like 'Google Meet', 'Zoom', 'Aether Office', etc."
                        }

                        Ensure accuracy in parsing and extracting details. If any field is missing or unclear, make a reasonable assumption.
                    '''
                },
                {'role': 'user', 'content': user_input}
            ],
            response_format=MeetingInfo
        )

        info = response.choices[0].message.parsed
        meeting_logger.info(f"LLM extraction successful - Title: '{info.title}', Date: '{info.date}', Time: '{info.time}'")

        if not info or not info.is_calendar_event or info.confidence_score < 0.6:
            meeting_logger.warning("No valid meeting detected or confidence too low")
            print("❌ No valid meeting detected.")
            return None
        
        meeting_logger.info("Meeting extraction completed successfully")
        # Final aggregated data
        return {
            "meeting_description":info.meeting_description,
            "is_calendar_event":info.is_calendar_event,
            "confidence_score":info.confidence_score,
            "title":info.title,
            "datetime":info.datetime,
            "date":info.date,
            "time":info.time,
            "duration_of_meeting":info.duration_of_meeting,
            "participants":info.participants,
            "confirmation_message":info.confirmation_message,
            "calendar_link":info.calendar_link,
            "notes":info.notes,
            "location":info.location
        }
    except Exception as e:
        meeting_logger.error(f"Error in meeting extraction process: {str(e)}")
        raise


# ======================================================================================================
#                                       DATA STORAGE CONFIGURATION in MongoDB
# ======================================================================================================

MEETINGS_FILE = "meeting_records"
RESUMES_FILE = "job_application"

myclient = pymongo.MongoClient("mongodb://localhost:27017/")
mydb = myclient["resumedb"]

def load_data(collection_name: str) -> List[Dict]:
    """Safely load data from MongoDB collection"""
    try:
        collection = mydb[collection_name]
        # Exclude MongoDB's _id field from results
        data = list(collection.find())
        streamlit_logger.debug(f"Successfully loaded {len(data)} records from {collection_name}")
        return data
    except Exception as e:
        streamlit_logger.error(f"Error loading data from MongoDB: {str(e)}")
        return []

def save_data(data: List[Dict], collection_name: str) -> None:
    """Safely save data to MongoDB collection"""
    try:
        collection = mydb[collection_name]
        
        result = collection.insert_one(data[-1])  # Use copy to avoid any modification issues
        streamlit_logger.info(f"Insert result: {result.inserted_id}")
            
        # Verify insertion
        if result.inserted_id:
            found = collection.find_one({"_id": result.inserted_id})
            if found:
                streamlit_logger.info("Document verified in database")
                return True
            else:
                streamlit_logger.error("Document not found after insertion")
                return False
        return False
        
    except Exception as e:
        streamlit_logger.error(f"Save error: {str(e)}", exc_info=True)
        return False

def serialize_pydantic_model(model) -> Dict:
    """Convert Pydantic model to serializable dict"""
    try:
        if isinstance(model, BaseModel):
            return model.model_dump()
        return model
    except Exception as e:
        streamlit_logger.error(f"Error serializing model: {str(e)}")
        raise ValueError(f"Could not serialize model: {str(e)}")


# ======================================================================================================
#                              Streamlit Application Implementation
# ======================================================================================================

st.set_page_config(page_title="Resume Parser", layout="wide")
st.sidebar.title("📄 Resume Parser App")


def save_parsed_resume(structured_data: dict, resume_text: str, filename: str) -> None:
    """Save parsed resume data to MongoDB"""
    try:
        global resume_jsn_data
        # Convert all Pydantic models to serializable dicts
        serializable_data = {
            section_name: serialize_pydantic_model(section_content)
            for section_name, section_content in structured_data.items()
        }
        
        # Create the complete record
        resume_record = {
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "parsed_data": serializable_data,
            "raw_text_sample": resume_text[:1000] + ("..." if len(resume_text) > 1000 else "")
        }
        # Save directly to MongoDB (no list handling needed)
        save_data([resume_record], "job_application")
        st.success("Resume data saved successfully to MongoDB!")
        
    except Exception as e:
        streamlit_logger.error(f"Failed to save resume data: {str(e)}")
        st.error(f"Failed to save resume data: {str(e)}")
        raise

def extract_text_from_pdf(file):
    try:
        streamlit_logger.info(f"Extracting text from PDF file: {file.name}")
        reader = PdfFileReader(file)
        text = ""
        for page_num, page in enumerate(reader.pages):
            streamlit_logger.debug(f"Processing page {page_num+1}")
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        streamlit_logger.info(f"Successfully extracted text from PDF file: {file.name}")
        return text.strip()
    except Exception as e:
        streamlit_logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_text_from_docx(file):
    try:
        streamlit_logger.info(f"Extracting text from DOCX file: {file.name}")
        doc = Document(file)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Also extract text from tables
        table_count = len(doc.tables)
        streamlit_logger.debug(f"Found {table_count} tables in document")
        for table_idx, table in enumerate(doc.tables):
            streamlit_logger.debug(f"Processing table {table_idx+1}")
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        
        streamlit_logger.info(f"Successfully extracted text from DOCX file: {file.name}")
        return "\n\n".join(paragraphs)
    except Exception as e:
        streamlit_logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

st.sidebar.title("Navigation")

# Define pages
pages = ["Resume Details", "Meeting Scheduler"]
selected_page = st.sidebar.radio("Navigate", pages)

if selected_page == "Resume Details":
    streamlit_logger.info("User navigated to Resume Details page")
    st.title("Resume Details")
    uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
    
    if uploaded_file is not None:
        file_extension = uploaded_file.name.split(".")[-1].lower()
        streamlit_logger.info(f"User uploaded file: {uploaded_file.name} ({file_extension})")
        
        if file_extension == "pdf":
            resume_text = extract_text_from_pdf(uploaded_file)
        elif file_extension == "docx":
            resume_text = extract_text_from_docx(uploaded_file)
        else:
            streamlit_logger.warning(f"Unsupported file format: {file_extension}")
            st.sidebar.error("Unsupported file format. Please upload a PDF or DOCX file.")
            st.stop()
        
        st.sidebar.subheader("Extracted Resume Text:")
        st.sidebar.text_area("Resume Content", resume_text, height=200)
        
        if st.sidebar.button("Parse Resume"):
            streamlit_logger.info("User clicked 'Parse Resume' button")
            try:
                structured_data = fn_Resume(resume_text)
                streamlit_logger.info("Successfully parsed resume")
                
                # Save with improved error handling
                try:
                    save_parsed_resume(structured_data, resume_text, uploaded_file.name)
                    st.success("Resume parsed and saved successfully!")
                except Exception as save_error:
                    st.error(f"Error saving resume: {str(save_error)}")
                    # Still show the parsed data even if saving failed
                
                # Display sections in tabs
                tabs = st.tabs(list(structured_data.keys()))
                
                for i, (tab_name, section_content) in enumerate(structured_data.items()):
                    with tabs[i]:
                        st.subheader(f"📌 {tab_name}")
                        if section_content:
                            display_data = serialize_pydantic_model(section_content)
                            st.json(display_data)
                        else:
                            st.write("No information available.")
                            
            except Exception as parse_error:
                streamlit_logger.error(f"Error parsing resume: {str(parse_error)}")
                st.error(f"Error parsing resume: {str(parse_error)}")
                st.stop()

        if st.sidebar.button("Download Resume PDF"):
            streamlit_logger.info("Generating pdf ")
            collection = mydb['job_application']
            data = list(collection.find())
            pdfmaker(data[-1])  # Ensure pdfmarker is correctly imported or aliased
            streamlit_logger.info("pdf generated")
            st.success("Resume PDF generated successfully!")
            with open('/home/ritik/Documents/vscode/AI_Agent/Resume_AI_Agent/resumes/professional_resume.pdf', "rb") as f:
                base64_pdf = base64.b64encode(f.read()).decode('utf-8')
            # Embedding PDF in HTML
            pdf_display = F'<embed src="data:application/pdf;base64,{base64_pdf}" width="700" height="1000" type="application/pdf">'

            # Displaying File
            st.markdown(pdf_display, unsafe_allow_html=True)


# ======================================================================================================
#                              Streamlit Application for Meeting Scheduler
# ======================================================================================================


else:

    def create_calendar_view(meetings):
        """Create a calendar-style view of meetings"""
        if not meetings:
            return st.info("No meetings scheduled yet.")
        
        # Get current month and year
        now = datetime.now()
        month = st.selectbox("Select Month", list(range(1, 13)), index=now.month-1)
        year = st.selectbox("Select Year", list(range(now.year-1, now.year+3)), index=1)
        
        # Create month calendar
        cal = calendar.monthcalendar(year, month)
        
        # Convert meetings to date-indexed dictionary
        meetings_by_date = {}
        for meeting in meetings:
            start_time = datetime.strptime(meeting['start'], "%Y-%m-%dT%H:%M:%S")
            if start_time.year == year and start_time.month == month:
                date_key = start_time.day
                if date_key not in meetings_by_date:
                    meetings_by_date[date_key] = []
                meetings_by_date[date_key].append(meeting)
        
        # Display calendar
        st.write(f"### {calendar.month_name[month]} {year}")
        
        # Create calendar grid
        cols = st.columns(7)
        for i, day_name in enumerate(["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]):
            cols[i].markdown(f"**{day_name}**")
        
        for week in cal:
            cols = st.columns(7)
            for i, day in enumerate(week):
                if day == 0:
                    cols[i].write("")
                else:
                    # Highlight today
                    if day == now.day and month == now.month and year == now.year:
                        cols[i].markdown(f"**{day}** 📅")
                    else:
                        cols[i].write(f"{day}")
                    
                    # List meetings for this day
                    if day in meetings_by_date:
                        for meeting in meetings_by_date[day]:
                            start_time = datetime.strptime(meeting['start'], "%Y-%m-%dT%H:%M:%S")
                            meeting_time = start_time.strftime("%H:%M")
                            cols[i].markdown(f"<small>🕒 {meeting_time} - {meeting['title']}</small>", unsafe_allow_html=True)

    streamlit_logger.info("User navigated to Meeting Scheduler page")
    # Streamlit UI
    st.title("📅 Meeting Scheduler")
    st.write("Extract meetings and visualize them on a calendar.")

    if "meetings" not in st.session_state:
        streamlit_logger.debug("Initializing meetings in session state")
        st.session_state.meetings = load_data(MEETINGS_FILE)

    # Create tabs for different views
    tab1, tab2, tab3 = st.tabs(["Add Meeting", "Calendar View", "Timeline View"])
    
    with tab1:
        user_input = st.text_area("📝 Enter your meeting request:", height=80)

        if st.button("🔍 Extract Meeting Details"):
            streamlit_logger.info("User clicked 'Extract Meeting Details' button")
            if user_input:
                streamlit_logger.info("Processing user input for meeting extraction")
                meeting_data = extract_meeting_info(user_input)
                if meeting_data:
                    streamlit_logger.info("Successfully extracted meeting details")
                    st.success("✅ Meeting details extracted successfully!")
                    st.json(meeting_data)
                    
                    try:
                        start_datetime = datetime.strptime(meeting_data['datetime'].replace("Z", ""), "%Y-%m-%dT%H:%M:%S")
                        end_datetime = start_datetime + timedelta(hours=meeting_data['duration_of_meeting'])
                        
                        meeting_event = {
                            **meeting_data,
                            "title": meeting_data['title'],
                            "start": meeting_data['datetime'],
                            "end": end_datetime.strftime("%Y-%m-%dT%H:%M:%S"),
                            "description": f"Participants: {', '.join(meeting_data['participants'])}"
                        }
                                                                        
                        if meeting_event not in st.session_state.meetings:
                            st.session_state.meetings.append(meeting_event)
                            save_data(st.session_state.meetings, MEETINGS_FILE)
                            st.success(f"Added: {meeting_data['title']} on {start_datetime.strftime('%Y-%m-%d')} at {start_datetime.strftime('%H:%M')}")
                            streamlit_logger.info(f"Added new meeting: {meeting_data['title']} at {start_datetime}")
                        else:
                            st.info("This meeting is already on your calendar.")
                            streamlit_logger.info(f"Skipped duplicate meeting: {meeting_data['title']}")
                    except Exception as e:
                        streamlit_logger.error(f"Error processing meeting datetime: {str(e)}")
                        st.error(f"❌ Error processing meeting datetime: {e}")
                else:
                    streamlit_logger.warning("No valid meeting detected in user input")
                    st.error("❌ No valid meeting detected. Please refine your input.")
            else:
                streamlit_logger.warning("User submitted empty meeting request")
                st.warning("⚠️ Please enter a meeting request.")

        # Manual meeting entry option
        with st.expander("Or add meeting manually"):
            with st.form("manual_meeting"):
                title = st.text_input("Meeting Title")
                date = st.date_input("Date")
                time = st.time_input("Time")
                duration = st.number_input("Duration (hours)", min_value=0.5, max_value=8.0, value=1.0, step=0.5)
                participants = st.text_input("Participants (comma-separated)")
                
                submit = st.form_submit_button("Add Meeting")
                if submit and title:
                    start_datetime = datetime.combine(date, time)
                    end_datetime = start_datetime + timedelta(hours=duration)
                    
                    meeting_event = {
                        "title": title,
                        "start": start_datetime.strftime("%Y-%m-%dT%H:%M:%S"),
                        "end": end_datetime.strftime("%Y-%m-%dT%H:%M:%S"),
                        "description": f"Participants: {participants}"
                    }
                    
                    if meeting_event not in st.session_state.meetings:
                        st.session_state.meetings.append(meeting_event)
                        save_data(st.session_state.meetings, MEETINGS_FILE)
                        st.success(f"Added: {title} on {date} at {time}")
                    else:
                        st.info("This meeting is already on your calendar.")
    
    with tab2:
        st.write(f"Total Meetings: {len(st.session_state.meetings)}")
        create_calendar_view(st.session_state.meetings)
    
    with tab3:
        if st.session_state.meetings:
            df = pd.DataFrame(st.session_state.meetings)
            df['start'] = pd.to_datetime(df['start'])
            df['end'] = pd.to_datetime(df['end'])
            df['duration'] = (df['end'] - df['start']).dt.total_seconds() / 3600
            df['status'] = df['start'].apply(lambda x: "Past" if x < datetime.now() else "Upcoming")
            
            fig = px.timeline(df, x_start="start", x_end="end", y="title", title="Meeting Schedule", color="status", 
                            text="title", color_discrete_map={"Past": "gray", "Upcoming": "green"})
            fig.update_layout(xaxis_title="Time", yaxis_title="Meetings", showlegend=True, height=400)
            fig.update_traces(textposition='inside')
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No meetings scheduled yet.")
    
    # Meeting list and management
    st.markdown("## 📋 Meeting List")
    
    if st.session_state.meetings:
        for i, event in enumerate(st.session_state.meetings):
            col1, col2 = st.columns([4, 1])
            start_time = datetime.strptime(event['start'], "%Y-%m-%dT%H:%M:%S")
            end_time = datetime.strptime(event['end'], "%Y-%m-%dT%H:%M:%S")
            status = "✅ Completed" if start_time < datetime.now() else "🕒 Upcoming"
            
            with col1:
                st.markdown(f"**{event['title']}** ({status})")
                st.markdown(f"📅 {start_time.strftime('%Y-%m-%d %H:%M')} - {end_time.strftime('%H:%M')}")
                st.markdown(f"👥 {event['description']}")
            
            with col2:
                if st.button("Delete", key=f"delete_{i}"):
                    st.session_state.meetings.pop(i)
                    save_data(st.session_state.meetings, MEETINGS_FILE)
                    st.success("Meeting deleted successfully!")
            
            st.divider()
    else:
        st.info("No meetings scheduled yet.")